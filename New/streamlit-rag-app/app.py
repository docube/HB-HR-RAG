import streamlit as st
from dotenv import load_dotenv
import os
from langchain.text_splitter import CharacterTextSplitter
from langchain.embeddings import OpenAIEmbeddings
from langchain.vectorstores import FAISS
from langchain.chat_models import ChatOpenAI
from langchain.chains import RetrievalQA
from langchain.docstore.document import Document
from PyPDF2 import PdfReader
import docx
import tempfile

# Load API Key
load_dotenv()
openai_api_key = os.getenv("OPENAI_API_KEY")

# Set up page config
st.set_page_config(page_title="HB-HR RAG App", layout="wide")
st.title("📂 HB-HR Retrieval-Augmented Generation (RAG) App")

# Initialize session state
if "vectorstore" not in st.session_state:
    st.session_state.vectorstore = None

# ---------- TAB 1: File Upload ----------
with st.sidebar:
    st.header("📁 Upload Files")
    uploaded_files = st.file_uploader(
        "Upload multiple documents",
        type=["pdf", "docx", "txt"],
        accept_multiple_files=True
    )

def read_file(file):
    if file.name.endswith(".pdf"):
        reader = PdfReader(file)
        return "\n".join([page.extract_text() or "" for page in reader.pages])
    elif file.name.endswith(".docx"):
        doc = docx.Document(file)
        return "\n".join([para.text for para in doc.paragraphs])
    elif file.name.endswith(".txt"):
        return file.read().decode("utf-8")
    return ""

tab1, tab2, tab3 = st.tabs(["📁 Upload", "🧠 Preprocess", "💬 Query"])

# ---------- TAB 1: Upload ----------
with tab1:
    st.subheader("Upload Documents")
    if uploaded_files:
        raw_texts = []
        for file in uploaded_files:
            text = read_file(file)
            raw_texts.append(text)
        st.session_state.raw_text = "\n".join(raw_texts)
        st.success(f"{len(uploaded_files)} document(s) uploaded and loaded!")

# ---------- TAB 2: Preprocessing & Vector Store ----------
with tab2:
    st.subheader("Preprocess & Embed")

    if st.button("🔍 Chunk + Embed"):
        if "raw_text" not in st.session_state:
            st.warning("Please upload documents first.")
        else:
            # Split into chunks
            text_splitter = CharacterTextSplitter(
                separator="\n",
                chunk_size=1000,
                chunk_overlap=200,
                length_function=len
            )
            docs = text_splitter.split_documents(
                [Document(page_content=st.session_state.raw_text)]
            )

            # Create vector store
            embeddings = OpenAIEmbeddings(openai_api_key=openai_api_key)
            vectorstore = FAISS.from_documents(docs, embeddings)
            st.session_state.vectorstore = vectorstore

            st.success("✅ Documents embedded successfully!")

# st.write("🔑 Loaded API Key (Preview 5 chars):", openai_api_key[:5] if openai_api_key else "❌ Not Loaded")


# ---------- TAB 3: Query ----------
with tab3:
    st.subheader("Ask Questions")

    if st.session_state.vectorstore is None:
        st.warning("Please embed your documents first.")
    else:
        query = st.text_input("Ask a question about your documents:")
        if query:
            llm = ChatOpenAI(temperature=0, openai_api_key=openai_api_key)
            qa_chain = RetrievalQA.from_chain_type(
                llm=llm,
                retriever=st.session_state.vectorstore.as_retriever()
            )
            result = qa_chain.run(query)
            st.write("📌 **Answer:**")
            st.info(result)

# st.write("🔑 Loaded API Key (Preview 5 chars):", openai_api_key[:5] if openai_api_key else "❌ Not Loaded")

